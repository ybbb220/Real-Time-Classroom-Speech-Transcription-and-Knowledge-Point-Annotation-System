from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database import get_db


def list_knowledge() -> list[dict]:
    # 查询全部知识点，按创建时间倒序
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, content, created_at FROM knowledge ORDER BY created_at DESC"
        )
        rows = cursor.fetchall()
        return [
            {"id": r[0], "content": r[1], "created_at": r[2].strftime("%Y-%m-%d %H:%M:%S")}
            for r in rows
        ]


def add_knowledge(content: str) -> int:
    # 新增一条知识点，返回新记录的id
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("INSERT INTO knowledge(content) VALUES(%s)", (content,))
        return cursor.lastrowid


def update_knowledge(knowledge_id: int, content: str) -> bool:
    # 更新知识点内容，返回是否成功
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE knowledge SET content=%s WHERE id=%s", (content, knowledge_id)
        )
        return cursor.rowcount > 0


def delete_knowledge(knowledge_id: int) -> bool:
    # 删除知识点，返回是否成功
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM knowledge WHERE id=%s", (knowledge_id,))
        return cursor.rowcount > 0


def get_all_docs() -> list[str]:
    # 返回所有知识点内容列表，供 FAISS 检索器建库使用
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT content FROM knowledge ORDER BY id")
        rows = cursor.fetchall()
        return [r[0] for r in rows]


def generate_docx(items: list[dict]) -> BytesIO:
    # items 格式: [{"text": "课堂语句", "doc": "知识点", "score": 0.85}, ...]
    doc = Document()

    title = doc.add_heading("知识点覆盖清单", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 空行

    table = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
    table.autofit = True

    headers = ["序号", "课堂语句", "匹配知识点"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)

    # 数据行
    for idx, item in enumerate(items, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = item.get("text", "")
        row_cells[2].text = item.get("doc", "")

        # 序号居中
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 统一字体大小
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)

    widths = [Cm(1.2), Cm(7), Cm(7)]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
